import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import pandas as pd
from docx import Document
from docx.shared import RGBColor
import os
import win32com.client
import threading
import pythoncom
import json
from PyPDF2 import PdfMerger

# --- Helper Functions ---
def sanitize_filename(name):
    invalid_chars = r'\/:*?"<>|'
    for ch in invalid_chars:
        name = name.replace(ch, '')
    return name.strip()

def upload_excel():
    global excel_file
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx;*.xls")])
    if file_path:
        excel_file = file_path
        lbl_excel.config(text=f"Excel File: {os.path.basename(file_path)}")
        if remember_paths_var.get():
            save_last_path('excel', file_path)

def upload_template():
    global template_file
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file_path:
        template_file = file_path
        lbl_template.config(text=f"Template File: {os.path.basename(file_path)}")
        if remember_paths_var.get():
            save_last_path('template', file_path)

def select_output_folder():
    global output_folder
    folder = filedialog.askdirectory()
    if folder:
        output_folder = folder
        lbl_output.config(text=f"Output Folder: {output_folder}")
        if remember_paths_var.get():
            save_last_path('output', folder)

def save_last_path(file_type, path):
    paths = {}
    if os.path.exists("last_paths.json"):
        with open("last_paths.json", "r") as f:
            paths = json.load(f)
    paths[file_type] = path
    with open("last_paths.json", "w") as f:
        json.dump(paths, f, indent=4)

def load_last_paths():
    global excel_file, template_file, output_folder
    if os.path.exists("last_paths.txt"):
        os.remove("last_paths.txt")
    if os.path.exists("last_paths.json"):
        try:
            with open("last_paths.json", "r") as f:
                paths = json.load(f)
            if 'excel' in paths:
                excel_file = paths['excel']
                lbl_excel.config(text=f"Excel File: {os.path.basename(excel_file)}")
            if 'template' in paths:
                template_file = paths['template']
                lbl_template.config(text=f"Template File: {os.path.basename(template_file)}")
            if 'output' in paths:
                output_folder = paths['output']
                lbl_output.config(text=f"Output Folder: {output_folder}")
        except Exception as e:
            print(f"Error loading paths: {e}")

def replace_placeholder_in_paragraph(paragraph, column_name, value):
    full_text = ''.join(run.text for run in paragraph.runs)
    placeholder = f"{{{{{column_name}}}}}"
    if placeholder not in full_text:
        return
    first_run_with_placeholder = None
    for run in paragraph.runs:
        if placeholder in run.text:
            first_run_with_placeholder = run
            break
    saved_font = None
    if first_run_with_placeholder:
        saved_font = first_run_with_placeholder.font
    for run in paragraph.runs:
        run.text = ""
    new_text = full_text.replace(placeholder, str(value))
    new_run = paragraph.add_run(new_text)
    if saved_font:
        font = new_run.font
        font.name = saved_font.name
        font.size = saved_font.size
        font.bold = saved_font.bold
        font.italic = saved_font.italic
        font.underline = saved_font.underline
    if column_name == "Car Number":
        font = new_run.font
        font.color.rgb = RGBColor(255, 0, 0)

def convert_docx_to_pdf(docx_path, pdf_path):
    try:
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"{docx_path} not found.")
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)
        doc.Close()
        word.Quit()
        pythoncom.CoUninitialize()
    except Exception as e:
        print(f"Error during PDF conversion: {e}")
        messagebox.showerror("Error", f"An error occurred during PDF conversion: {e}")

def merge_pdfs(pdf_list, output_pdf):
    try:
        pdf_merger = PdfMerger()
        for pdf in pdf_list:
            pdf_merger.append(pdf)
        pdf_merger.write(output_pdf)
        pdf_merger.close()
    except Exception as e:
        print(f"Error during merging PDFs: {e}")
        messagebox.showerror("Error", f"An error occurred during merging PDFs: {e}")

# --- Global Flags ---
cancel_flag = False
pause_flag = False
pause_event = threading.Event()
pause_event.set()

# List to store the generated PDFs
generated_pdfs = []

# --- Main PDF Generation ---
def generate_pdfs():
    global cancel_flag, pause_flag, generated_pdfs
    if not excel_file or not template_file or not output_folder:
        messagebox.showerror("Error", "Please upload Excel, Word template, and select an output folder.")
        return
    if cancel_flag:
        cancel_flag = False
        progress['value'] = 0
        lbl_status.config(text="Status: Ready")
    try:
        lbl_status.config(text="Generating PDFs...")
        root.update_idletasks()
        df = pd.read_excel(excel_file)
        os.makedirs(output_folder, exist_ok=True)
        column_mapping = {
            "Name": "Namn",
            "Address": "Adress",
            "Postal Code": "Postadress",
            "Postal number": "Postnr",
            "Car Number": "Registreringsnr"
        }
        progress['maximum'] = len(df)
        progress['value'] = 0
        for index, row in df.iterrows():
            if cancel_flag:
                lbl_status.config(text="Process cancelled.")
                messagebox.showinfo("Cancelled", "PDF generation has been cancelled.")
                # Merge the PDFs generated so far and save it
                if generated_pdfs:
                    merged_pdf_path = os.path.join(output_folder, "All_Reklam.pdf")
                    merge_pdfs(generated_pdfs, merged_pdf_path)
                    messagebox.showinfo("Success", f"Reklams merged into one PDF!\nSaved in:\n{merged_pdf_path}")
                return
            pause_event.wait()
            temp_doc = Document(template_file)
            for placeholder, column_name in column_mapping.items():
                value = row.get(column_name, '')
                if isinstance(value, float) and value.is_integer():
                    value = str(int(value))
                else:
                    value = str(value)
                for paragraph in temp_doc.paragraphs:
                    replace_placeholder_in_paragraph(paragraph, placeholder, value)
                for table in temp_doc.tables:
                    for row_cells in table.rows:
                        for cell in row_cells.cells:
                            for paragraph in cell.paragraphs:
                                replace_placeholder_in_paragraph(paragraph, placeholder, value)
            temp_doc_filename = os.path.join(output_folder, f"Reklam_{index + 1}_modified.docx")
            temp_doc.save(temp_doc_filename)
            pdf_filename = os.path.join(output_folder, f"Reklam_{index + 1}.pdf")
            convert_docx_to_pdf(temp_doc_filename, pdf_filename)
            os.remove(temp_doc_filename)
            generated_pdfs.append(pdf_filename)  # Add the generated PDF to the list
            progress['value'] += 1
            root.update_idletasks()
        # After all PDFs are generated, merge them
        if generated_pdfs:
            merged_pdf_path = os.path.join(output_folder, "All_Reklam.pdf")
            merge_pdfs(generated_pdfs, merged_pdf_path)
            messagebox.showinfo("Success", f"Reklams merged into one PDF!\nSaved in:\n{merged_pdf_path}")
        lbl_status.config(text="PDFs generated and merged successfully!")
    except Exception as e:
        lbl_status.config(text="Error during generation.")
        messagebox.showerror("Error", f"An error occurred:\n{e}")

# --- Control Functions for Buttons ---
def cancel_generation():
    global cancel_flag
    cancel_flag = True
    pause_event.set()
    lbl_status.config(text="Cancelled")

def pause_generation():
    global pause_flag
    pause_flag = True
    pause_event.clear()
    lbl_status.config(text="Paused...")

def resume_generation():
    global pause_flag
    pause_flag = False
    pause_event.set()
    lbl_status.config(text="Resuming...")

# --- GUI Setup ---
root = tk.Tk()
root.title("Excel to PDF Reklam Generator")
root.geometry("800x650")
root.resizable(True, True)
root.configure(bg="#f5f5f5")

excel_file = ""
template_file = ""
output_folder = ""

# Section: Upload Files
section_upload = tk.Label(root, text="1. Upload Files", font=("Arial", 14, "bold"), bg="#f5f5f5")
section_upload.pack(pady=(10, 5))

frame_upload = tk.Frame(root, bg="#f5f5f5")
frame_upload.pack(pady=5)

btn_excel = tk.Button(frame_upload, text="Upload Excel File", command=upload_excel, width=25)
btn_excel.grid(row=0, column=0, padx=10, pady=5)

lbl_excel = tk.Label(frame_upload, text="No Excel file selected", bg="#f5f5f5")
lbl_excel.grid(row=0, column=1, padx=10, pady=5)

btn_template = tk.Button(frame_upload, text="Upload Word Template", command=upload_template, width=25)
btn_template.grid(row=1, column=0, padx=10, pady=5)

lbl_template = tk.Label(frame_upload, text="No Word template selected", bg="#f5f5f5")
lbl_template.grid(row=1, column=1, padx=10, pady=5)

# Section: Output Folder
section_output = tk.Label(root, text="2. Select Output Folder", font=("Arial", 14, "bold"), bg="#f5f5f5")
section_output.pack(pady=(20, 5))

frame_output = tk.Frame(root, bg="#f5f5f5")
frame_output.pack(pady=5)

btn_output = tk.Button(frame_output, text="Select Output Folder", command=select_output_folder, width=25)
btn_output.grid(row=0, column=0, padx=10, pady=5)

lbl_output = tk.Label(frame_output, text="No output folder selected", bg="#f5f5f5")
lbl_output.grid(row=0, column=1, padx=10, pady=5)

# Section: Remember Last Paths
remember_paths_var = tk.IntVar()
chk_remember_paths = tk.Checkbutton(root, text="Remember Last Used Files/Folders", variable=remember_paths_var, bg="#f5f5f5")
chk_remember_paths.pack(pady=10)

btn_generate = tk.Button(root, text="Start Generating Reklams", command=lambda: threading.Thread(target=generate_pdfs).start(), width=30, bg="#4CAF50", fg="white", font=("Arial", 12, "bold"))
btn_generate.pack(pady=10)

# Section: Progress
progress = ttk.Progressbar(root, orient="horizontal", length=400, mode="determinate")
progress.pack(pady=15)

lbl_status = tk.Label(root, text="Status: Ready", bg="#f5f5f5")
lbl_status.pack(pady=5)

# Section: Control Buttons
frame_controls = tk.Frame(root, bg="#f5f5f5")
frame_controls.pack(pady=10)

# Increase the width to make the buttons bigger
btn_cancel = tk.Button(frame_controls, text="Cancel", command=cancel_generation, width=20, bg="#f44336", fg="white", font=("Arial", 12, "bold"))
btn_cancel.grid(row=0, column=0, padx=10)

btn_pause = tk.Button(frame_controls, text="Pause", command=pause_generation, width=20, bg="#ff9800", fg="white", font=("Arial", 12, "bold"))
btn_pause.grid(row=0, column=1, padx=10)

btn_resume = tk.Button(frame_controls, text="Resume", command=resume_generation, width=20, bg="#4CAF50", fg="white", font=("Arial", 12, "bold"))
btn_resume.grid(row=0, column=2, padx=10)


# Load last used paths if available
load_last_paths()

root.mainloop()
